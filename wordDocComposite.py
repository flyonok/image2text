from docx import Document


def CompositeTwoDocs(srcDocFullName, dstDocFullName, compositeName):
    '''
    srcDocFullName:源文档，里面含有需要替换的内容
    dstDocFullName:目标文档，执行后，相关模板内容被替换
    compositeName:替换的对象名，比如正面或背面
    return: 成功->True，失败->False
    '''
    try:
        srcDoc = Document(srcDocFullName)
        dstDoc = Document(dstDocFullName)
        srcParasMap = {}  # Heading 2 => [paras list]
        dstParasMap = {}  # Heading 2 => [paras list]
        firstPage = False
        secondPage = False
        currentLabelStyleContent = None  # 当前标签样式对应的内容
        # 查找源文档的相关内容
        for srcPara in srcDoc.paragraphs:
            if (srcPara.style.name.find('Heading 2') >= 0 and srcPara.text.find(compositeName) >= 0):
                print('find {0}'.format(srcPara))
                firstPage = True
            elif (srcPara.style.name.find('Heading 2') >= 0 and firstPage):
                secondPage = True
                break
            else:
                if (firstPage and not secondPage):
                    if (srcPara.style.name.find('Heading 3') >= 0):
                        srcParasMap[srcPara.text] = []
                        currentLabelStyleContent = srcPara.text
                    else:
                        if currentLabelStyleContent is None:
                            raise ValueError('不合格的word模板文档！')
                        srcParasMap[currentLabelStyleContent].append(srcPara)
        firstPage = False
        secondPage = False
        currentLabelStyleContent = None  # 当前标签样式对应的内容
        # 查找目标文档的相关内容
        for dstPara in dstDoc.paragraphs:
            if (dstPara.style.name.find('Heading 2') >= 0 and dstPara.text.find(compositeName) >= 0):
                print('find {0}'.format(dstPara))
                firstPage = True
            elif (dstPara.style.name.find('Heading 2') >= 0 and firstPage):
                secondPage = True
                break
            else:
                if (firstPage and not secondPage):
                    if (dstPara.style.name.find('Heading 3') >= 0):
                        dstParasMap[dstPara.text] = []
                        currentLabelStyleContent = dstPara.text
                    else:
                        if currentLabelStyleContent is None:
                            raise ValueError('不合格的word模板文档！')
                        dstParasMap[currentLabelStyleContent].append(dstPara)
        
        # 开始组合
        for key, dstParas in dstParasMap.items():
            srcParas = srcParasMap[key]
            if len(srcParas) <= 0:
                print('源文档中没有该项--{0}--内容'.format(key))
                continue
            else:
                for index, item in enumerate(dstParas):
                    if (index <= len(srcParas)):
                        dstParas[index].text = srcParas[index].text
                    else:
                        print('{0}中的长度--{1}--已经大于源文档的总长度--{2}'.format(key, index, len(srcParas)))
        dstDoc.save(dstDocFullName)

    except Exception as e:
        print('出现错误...')
        print(e)
        return False
    return True

if __name__ == '__main__':
    srcDocFullName = r'D:\秒秒学人工智能平台\2020年8月\名片-111\名片-111.docx'
    dstDocFullName = r'D:\秒秒学人工智能平台\2020年8月\名片-456\名片-456.docx'
    CompositeTwoDocs(srcDocFullName, dstDocFullName, '正面')
    